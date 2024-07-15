# This code writes down sequential call numbers into a given template. 
# Simply upload the template named "input.docx" into this folder, 
# edit the start_value and end_value with the first number you want on the page
# and the last numebr on the page. 

# input the first and last call numbers you want
start_value = 16760
end_value = 16904

from docx import Document
from docx.enum.text import WD_BREAK
import re
from docx.shared import Pt
from docx.oxml.ns import qn

def add_page_break(paragraph):
    paragraph.add_run().add_break(WD_BREAK.PAGE)

def save_partial_doc(doc, output_path, start_num, end_num):
    doc.save(f"{output_path}_{start_num}_to_{end_num}.docx")

def update_numbers_in_docx(file_path, output_path, start_value, end_value):
    numbers_per_page = 30
    max_numbers_per_file = 80

    def update_numbers(text, current_value):
        def replacement(match):
            nonlocal current_value
            updated_number = str(current_value)
            current_value += 1
            return updated_number
        new_text = re.sub(r'\b\d+\b', replacement, text)
        return new_text, current_value

    def format_run(run):
        run.bold = True
        run.font.name = 'Arial'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        run.font.size = Pt(13)

    def process_paragraphs(paragraphs, current_value):
        number_count = 0
        for paragraph in paragraphs:
            new_text, current_value = update_numbers(paragraph.text, current_value)
            if new_text != paragraph.text:
                paragraph.clear()
                run = paragraph.add_run(new_text)
                format_run(run)
            number_count += len(re.findall(r'\b\d+\b', paragraph.text))
            if number_count >= numbers_per_page:
                add_page_break(paragraph)
                number_count = 0
        return current_value

    current_start = start_value

    while current_start <= end_value:
        current_end = min(current_start + max_numbers_per_file - 1, end_value)
        current_value = current_start

        doc = Document(file_path)
        current_value = process_paragraphs(doc.paragraphs, current_value)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    current_value = process_paragraphs(cell.paragraphs, current_value)

        output_file = f"{output_path}_{current_start}_to_{current_end}.docx"
        save_partial_doc(doc, output_file, current_start, current_value - 1)

        current_start = current_end + 1

# Example usage
input_file = 'template labeling/input.docx'
output_path = 'template labeling/output/'

update_numbers_in_docx(input_file, output_path, start_value, end_value)


