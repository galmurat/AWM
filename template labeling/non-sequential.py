# If your call numbers are not sequential (has multiple parts
# you can copy and paste your call numbers here (seperate with commas and quotation marks).
# Your output will appear in the output folder. 

from docx import Document
from docx.enum.text import WD_BREAK
import re
from docx.shared import Pt
from docx.oxml.ns import qn

# Replace with values you want. 
values = [
"AWM RL 16750","AWM RL 16751","AWM RL 16751","AWM RL 16751","AWM RL 16752","AWM RL 16752","AWM RL 16752","AWM RL 16752",
"AWM RL 16752","AWM RL 16752","AWM RL 16752","AWM RL 16752","AWM RL 16752","AWM RL 16752","AWM RL 16752","AWM RL 16752",
"AWM RL 16752","AWM RL 16752","AWM RL 16752","AWM RL 16752","AWM RL 16753","AWM RL 16753","AWM RL 16753","AWM RL 16753",
"AWM RL 16753","AWM RL 16753","AWM RL 16753","AWM RL 16753","AWM RL 16753","AWM RL 16753","AWM RL 16753","AWM RL 16753",
"AWM RL 16753","AWM RL 16753","AWM RL 16753","AWM RL 16753","AWM RL 16753","AWM RL 16754","AWM RL 16754","AWM RL 16754",
"AWM RL 16754","AWM RL 16754","AWM RL 16754","AWM RL 16754","AWM RL 16754","AWM RL 16754","AWM RL 16754","AWM RL 16754",
"AWM RL 16755","AWM RL 16755","AWM RL 16755","AWM RL 16755","AWM RL 16755","AWM RL 16755","AWM RL 16755","AWM RL 16755",
"AWM RL 16755","AWM RL 16755","AWM RL 16755","AWM RL 16755","AWM RL 16755","AWM RL 16755","AWM RL 16756","AWM RL 16756",
"AWM RL 16756","AWM RL 16756","AWM RL 16756","AWM RL 16756","AWM RL 16756","AWM RL 16756","AWM RL 16756","AWM RL 16756",
"AWM RL 16756","AWM RL 16756","AWM RL 16757","AWM RL 16757","AWM RL 16757","AWM RL 16757","AWM RL 16757","AWM RL 16758",
"AWM RL 16759"
]


def add_page_break(paragraph):
    paragraph.add_run().add_break(WD_BREAK.PAGE)

def save_partial_doc(doc, output_path, start_num, end_num):
    doc.save(f"{output_path}_{start_num}_to_{end_num}.docx")

def update_numbers_in_docx(file_path, output_path, values):
    numbers_per_page = 30
    
    max_numbers_per_file = 80

    def update_numbers(text, values, index):
        pattern = r'AWM SC \d+(\(\d+\))?'  # Adjust the pattern to match the full text
        def replacement(match):
            nonlocal index
            if index < len(values):
                updated_number = values[index]
                index += 1
                return updated_number
            else:
                return match.group(0)
        new_text = re.sub(pattern, replacement, text)
        return new_text, index

    def format_run(run):
        run.bold = True
        run.font.name = 'Arial'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        run.font.size = Pt(11)

    def process_paragraphs(paragraphs, values, index):
        number_count = 0
        for paragraph in paragraphs:
            new_text, index = update_numbers(paragraph.text, values, index)
            if new_text != paragraph.text:
                paragraph.clear()
                run = paragraph.add_run(new_text)
                format_run(run)
            number_count += len(re.findall(r'AWM SC \d+(\(\d+\))?', paragraph.text))
            if number_count >= numbers_per_page:
                add_page_break(paragraph)
                number_count = 0
        return index

    total_values = len(values)
    current_start = 0

    while current_start < total_values:
        current_end = min(current_start + max_numbers_per_file, total_values)
        index = current_start

        doc = Document(file_path)
        index = process_paragraphs(doc.paragraphs, values, index)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    index = process_paragraphs(cell.paragraphs, values, index)

        output_file = f"{output_path}_{values[current_start]}_to_{values[current_end - 1]}.docx"
        save_partial_doc(doc, output_file, values[current_start], values[current_end - 1])

        current_start = current_end
# Example usage
input_file = 'template labeling/input2.docx'
output_path = 'template labeling/output/'

update_numbers_in_docx(input_file, output_path, values)

